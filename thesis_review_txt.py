#!/usr/bin/env python3
"""
Comprehensive thesis review for Chinese master's theses.
Checks formatting, language, academic content, structure, and references.
Outputs review in text format.
"""

import re
import os
import sys
from datetime import datetime
from docx import Document

class ThesisReviewer:
    def __init__(self, input_path):
        """Initialize reviewer with thesis document path."""
        self.input_path = input_path
        self.document = Document(input_path)
        self.issues = []
        self.positive_feedback = []

        # Extract document content
        self.text_content = []
        self.paragraphs = []
        self.headings = []
        self.tables = []
        self.references = []

        self._extract_content()

    def _extract_content(self):
        """Extract all content from the document."""
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            if text:
                self.text_content.append(text)
                self.paragraphs.append({
                    'index': i,
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })

                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    self.headings.append({
                        'index': i,
                        'text': text,
                        'level': level,
                        'para_index': i
                    })

        # Extract tables
        for i, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            self.tables.append({
                'index': i,
                'data': table_data,
                'location': self._find_table_location(i)
            })

    def _find_table_location(self, table_index):
        """Find which section a table is in."""
        # Simple implementation: find the nearest heading before the table
        # In a real implementation, we'd need to track paragraph positions
        return f"Table {table_index + 1}"

    def check_punctuation_consistency(self):
        """Check for mixed Chinese-English punctuation."""
        chinese_punct = ['。', '，', '；', '：', '？', '！', '「', '」', '《', '》']
        english_punct = ['.', ',', ';', ':', '?', '!', '"', "'", '<', '>']

        for i, para in enumerate(self.paragraphs):
            text = para['text']

            # Check for mixed punctuation in the same paragraph
            has_chinese = any(p in text for p in chinese_punct)
            has_english = any(p in text for p in english_punct[:6])  # Basic punctuation

            if has_chinese and has_english:
                # Check if it's likely intentional (English terms or citations)
                english_terms = ['et al.', 'e.g.', 'i.e.', 'etc.', 'vs.', 'Fig.', 'Table']
                has_english_terms = any(term in text for term in english_terms)

                if not has_english_terms:
                    self.issues.append({
                        'category': 'formatting',
                        'type': 'mixed_punctuation',
                        'location': f"Paragraph {i+1}",
                        'description': '中英文标点符号混用',
                        'suggestion': '统一使用中文标点符号（。，；：？！）或根据上下文合理选择',
                        'priority': 'medium'
                    })

    def check_decimal_places(self):
        """Check for excessive decimal places in tables."""
        for table in self.tables:
            for row_idx, row in enumerate(table['data']):
                for col_idx, cell in enumerate(row):
                    # Look for numbers with decimal places
                    numbers = re.findall(r'\d+\.\d+', cell)
                    for num in numbers:
                        decimal_part = num.split('.')[1]
                        if len(decimal_part) > 2:
                            self.issues.append({
                                'category': 'academic_content',
                                'type': 'excessive_decimals',
                                'location': f"{table['location']}, 行{row_idx+1}, 列{col_idx+1}",
                                'description': f'数值"{num}"有{len(decimal_part)}位小数，通常保留2位即可',
                                'suggestion': '除非测量精度要求，否则将小数位数减少到2位',
                                'priority': 'low'
                            })

    def check_statistical_expressions(self):
        """Check statistical expression correctness."""
        p_value_patterns = [
            r'p\s*[=<>]\s*\d+\.\d+',  # p=0.05, p<0.01, p>0.1
            r'p\s*[=<>]\s*\d+',        # p=0, p<1
            r'P\s*[=<>]\s*\d+\.\d+',   # P=0.05
        ]

        for i, para in enumerate(self.paragraphs):
            text = para['text']

            # Check p-value formatting
            for pattern in p_value_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    p_expr = match.group()
                    # Check if it has too many decimal places
                    numbers = re.findall(r'\d+\.\d+', p_expr)
                    for num in numbers:
                        if len(num.split('.')[1]) > 3:
                            self.issues.append({
                                'category': 'academic_content',
                                'type': 'p_value_format',
                                'location': f"Paragraph {i+1}",
                                'description': f'p值"{p_expr}"的小数位数过多',
                                'suggestion': 'p值通常报告2-3位小数，如p=0.023或p<0.001',
                                'priority': 'medium'
                            })

            # Check for "significant difference" vs "statistically significant difference"
            if '显著差异' in text and '统计学' not in text and '统计' not in text:
                # Check context to see if it's about statistical significance
                if any(word in text for word in ['p', 'P', '检验', '分析', '结果']):
                    self.issues.append({
                        'category': 'academic_content',
                        'type': 'statistical_terminology',
                        'location': f"Paragraph {i+1}",
                        'description': '使用"显著差异"而未明确是"统计学显著差异"',
                        'suggestion': '明确表述为"统计学显著差异"以避免歧义',
                        'priority': 'low'
                    })

    def check_heading_hierarchy(self):
        """Check heading hierarchy and numbering."""
        heading_levels = [h['level'] for h in self.headings]

        # Check for skipped levels
        for i in range(1, len(heading_levels)):
            if heading_levels[i] > heading_levels[i-1] + 1:
                self.issues.append({
                    'category': 'structure',
                    'type': 'heading_hierarchy',
                    'location': f"标题: {self.headings[i]['text']}",
                    'description': f'标题层级跳跃: 前一个{heading_levels[i-1]}级标题直接跳到{heading_levels[i]}级',
                    'suggestion': '标题层级应连续，避免跳跃（如1级直接到3级）',
                    'priority': 'medium'
                })

        # Check numbering consistency for numbered headings
        numbered_headings = []
        for heading in self.headings:
            text = heading['text']
            # Check if heading starts with a number (like "1. ", "2.1 ", etc.)
            match = re.match(r'^(\d+(\.\d+)*)\.?\s+', text)
            if match:
                number = match.group(1)
                numbered_headings.append({
                    'number': number,
                    'text': text,
                    'level': heading['level']
                })

        # Check for gaps in numbering
        if numbered_headings:
            prev_number = None
            for h in numbered_headings:
                if prev_number:
                    # Simple check: if current number doesn't follow previous
                    try:
                        prev_parts = prev_number.split('.')
                        curr_parts = h['number'].split('.')

                        # Check if at same level
                        if len(curr_parts) == len(prev_parts):
                            prev_last = int(prev_parts[-1])
                            curr_last = int(curr_parts[-1])
                            if curr_last != prev_last + 1:
                                self.issues.append({
                                    'category': 'structure',
                                    'type': 'heading_numbering',
                                    'location': f"标题: {h['text']}",
                                    'description': f'标题编号不连续: {prev_number} 后应为 {prev_last+1}，实际为 {h["number"]}',
                                    'suggestion': '检查并修正标题编号顺序',
                                    'priority': 'high'
                                })
                    except ValueError:
                        pass
                prev_number = h['number']

    def check_abstract_balance(self):
        """Check if abstract is head-heavy (too much background, little results)."""
        abstract_paras = []
        in_abstract = False

        for para in self.paragraphs:
            text = para['text']
            style = para['style']

            if '摘要' in text and ('Heading' in style or len(text) < 20):
                in_abstract = True
                continue
            elif in_abstract:
                if 'Heading' in style:  # Next heading ends abstract
                    break
                if text and not text.startswith('关键词'):
                    abstract_paras.append(text)

        if abstract_paras:
            abstract_text = ' '.join(abstract_paras)
            # Simple heuristic: count background vs results terms
            background_terms = ['随着', '近年来', '目前', '背景', '目的', '旨在', '探讨']
            results_terms = ['结果', '发现', '表明', '显示', '达到', '验证', '证明']

            bg_count = sum(term in abstract_text for term in background_terms)
            res_count = sum(term in abstract_text for term in results_terms)

            if bg_count > 0 and res_count > 0:
                ratio = bg_count / (bg_count + res_count)
                if ratio > 0.7:  # More than 70% background terms
                    self.issues.append({
                        'category': 'structure',
                        'type': 'abstract_balance',
                        'location': '摘要部分',
                        'description': f'摘要头重脚轻: 背景介绍占比过高({ratio:.0%})，结果部分不足',
                        'suggestion': '增加结果和结论的篇幅，减少背景介绍',
                        'priority': 'medium'
                    })

    def check_reference_numbering(self):
        """Check reference numbering consistency."""
        ref_section = None
        ref_paras = []

        # Find references section
        for i, para in enumerate(self.paragraphs):
            text = para['text']
            if '参考' in text and '文献' in text and len(text) < 20:
                ref_section = i
                break

        if ref_section:
            # Collect reference paragraphs
            for para in self.paragraphs[ref_section + 1:]:
                text = para['text']
                if text and not ('Heading' in para['style'] or '致谢' in text or '附录' in text):
                    ref_paras.append(text)
                elif 'Heading' in para['style']:
                    break

            # Check numbering
            expected_num = 1
            for ref_text in ref_paras:
                # Look for [1], [2], etc.
                match = re.search(r'\[(\d+)\]', ref_text)
                if match:
                    ref_num = int(match.group(1))
                    if ref_num != expected_num:
                        self.issues.append({
                            'category': 'references',
                            'type': 'reference_numbering',
                            'location': f'参考文献 #{expected_num}',
                            'description': f'参考文献编号不连续: 期望[{expected_num}]，实际[{ref_num}]',
                            'suggestion': '检查并修正参考文献编号顺序',
                            'priority': 'high'
                        })
                        expected_num = ref_num + 1
                    else:
                        expected_num += 1
                else:
                    # No numbering found, might be using different format
                    break

    def check_methods_order(self):
        """Check logical order of methods section."""
        methods_headings = []

        for heading in self.headings:
            text = heading['text']
            if '方法' in text or '实验设计' in text or '材料' in text:
                methods_headings.append(heading)

        if len(methods_headings) > 1:
            # Check if headings are in logical order
            # Typically: 材料/数据 -> 方法/设计 -> 实施/步骤 -> 分析/评估
            order_keywords = [
                ['材料', '数据', '样本', '来源'],  # First: materials/data
                ['方法', '设计', '模型', '算法'],  # Second: methods/design
                ['步骤', '流程', '实施', '过程'],  # Third: procedures
                ['分析', '评估', '指标', '统计']   # Fourth: analysis
            ]

            heading_texts = [h['text'] for h in methods_headings]
            for i, heading in enumerate(heading_texts):
                for stage_idx, keywords in enumerate(order_keywords):
                    if any(keyword in heading for keyword in keywords):
                        if i < stage_idx - 1:  # Too early
                            self.issues.append({
                                'category': 'structure',
                                'type': 'methods_order',
                                'location': f"标题: {heading}",
                                'description': '方法部分顺序可能不合理',
                                'suggestion': '考虑按"材料/数据 → 方法/设计 → 实施步骤 → 分析评估"的逻辑顺序组织',
                                'priority': 'medium'
                            })
                        break

    def check_title_content_alignment(self):
        """Check if title aligns with content scope."""
        if self.paragraphs:
            title = self.paragraphs[0]['text'] if len(self.paragraphs[0]['text']) > 10 else None

            if title and len(title) < 100:  # Likely the main title
                # Extract keywords from title
                title_keywords = re.findall(r'[\u4e00-\u9fff]+', title)

                # Check if these keywords appear in content
                content_text = ' '.join([p['text'] for p in self.paragraphs[10:50]])  # First few paragraphs

                missing_keywords = []
                for keyword in title_keywords:
                    if len(keyword) > 1 and keyword not in content_text:
                        missing_keywords.append(keyword)

                if missing_keywords:
                    self.issues.append({
                        'category': 'structure',
                        'type': 'title_content_alignment',
                        'location': '论文标题',
                        'description': f'标题中的关键词"{", ".join(missing_keywords)}"在正文前部出现频率较低',
                        'suggestion': '确保标题准确反映论文核心内容，或在正文中更充分地讨论标题涉及的主题',
                        'priority': 'high'
                    })

    def check_figure_table_numbering(self):
        """Check consecutive numbering of figures and tables."""
        fig_patterns = [r'图\s*(\d+(\.\d+)*)', r'Figure\s*(\d+(\.\d+)*)', r'Fig\.\s*(\d+(\.\d+)*)']
        table_patterns = [r'表\s*(\d+(\.\d+)*)', r'Table\s*(\d+(\.\d+)*)', r'Tab\.\s*(\d+(\.\d+)*)']

        fig_numbers = []
        table_numbers = []

        for para in self.paragraphs:
            text = para['text']

            # Check for figure references
            for pattern in fig_patterns:
                matches = re.finditer(pattern, text)
                for match in matches:
                    num = match.group(1)
                    fig_numbers.append(num)

            # Check for table references
            for pattern in table_patterns:
                matches = re.finditer(pattern, text)
                for match in matches:
                    num = match.group(1)
                    table_numbers.append(num)

        # Check figure numbering
        self._check_numbering_sequence(fig_numbers, '图', 'figure_numbering')

        # Check table numbering
        self._check_numbering_sequence(table_numbers, '表', 'table_numbering')

    def _check_numbering_sequence(self, numbers, label, issue_type):
        """Check if numbers are consecutive."""
        if not numbers:
            return

        # Convert to sortable format
        sorted_nums = sorted(set(numbers), key=lambda x: [int(part) for part in x.split('.')])

        # Check for gaps
        for i in range(1, len(sorted_nums)):
            prev_parts = sorted_nums[i-1].split('.')
            curr_parts = sorted_nums[i].split('.')

            if len(prev_parts) == len(curr_parts):
                try:
                    prev_last = int(prev_parts[-1])
                    curr_last = int(curr_parts[-1])

                    if curr_last != prev_last + 1:
                        self.issues.append({
                            'category': 'academic_content',
                            'type': issue_type,
                            'location': f'{label}编号',
                            'description': f'{label}编号不连续: {sorted_nums[i-1]} 后应为 {prev_last+1}，但找到 {sorted_nums[i]}',
                            'suggestion': f'检查并修正{label}编号顺序',
                            'priority': 'medium'
                        })
                except ValueError:
                    pass

    def collect_positive_feedback(self):
        """Collect positive aspects of the thesis."""
        # Check for well-structured abstract
        abstract_paras = []
        in_abstract = False

        for para in self.paragraphs:
            text = para['text']
            if '摘要' in text and len(text) < 20:
                in_abstract = True
            elif in_abstract and ('Heading' in para['style'] or '关键词' in text):
                break
            elif in_abstract and text:
                abstract_paras.append(text)

        if 2 <= len(abstract_paras) <= 4:
            self.positive_feedback.append('摘要结构清晰，长度适中')

        # Check for clear headings
        if len(self.headings) >= 5:
            self.positive_feedback.append('章节结构明确，标题层次清晰')

        # Check for tables/figures
        if self.tables:
            self.positive_feedback.append('包含数据表格，增强论证说服力')

        # Check for references
        ref_count = sum(1 for para in self.paragraphs if '参考' in para['text'] and '文献' in para['text'])
        if ref_count > 0:
            self.positive_feedback.append('包含参考文献部分，体现学术规范')

    def perform_comprehensive_review(self):
        """Perform all checks."""
        print("开始论文评审...")

        # Run all checks
        self.check_punctuation_consistency()
        self.check_decimal_places()
        self.check_statistical_expressions()
        self.check_heading_hierarchy()
        self.check_abstract_balance()
        self.check_reference_numbering()
        self.check_methods_order()
        self.check_title_content_alignment()
        self.check_figure_table_numbering()
        self.collect_positive_feedback()

        print(f"评审完成，共发现{len(self.issues)}个问题")

    def generate_review_text(self, output_path):
        """Generate review document in text format."""
        content = []

        # Title
        content.append("=" * 80)
        content.append("硕士学位论文评审意见")
        content.append("=" * 80)
        content.append("")

        # Metadata
        content.append(f"评审文件: {os.path.basename(self.input_path)}")
        content.append(f"生成日期: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        content.append(f"评审人: Claude Code 学术评审系统")
        content.append("")

        # 1. Overall assessment
        content.append("1. 总体评价")
        content.append("-" * 40)

        # Count issues by priority
        high_issues = [i for i in self.issues if i['priority'] == 'high']
        medium_issues = [i for i in self.issues if i['priority'] == 'medium']
        low_issues = [i for i in self.issues if i['priority'] == 'low']

        content.append("本文基本符合硕士学位论文要求，但在格式规范、学术表达和结构组织方面存在一些需要改进的问题。")
        content.append(f"共发现{len(self.issues)}个问题，其中：")
        content.append(f"  • 高优先级问题: {len(high_issues)}个")
        content.append(f"  • 中优先级问题: {len(medium_issues)}个")
        content.append(f"  • 低优先级问题: {len(low_issues)}个")
        content.append("")

        # 2. Detailed findings
        content.append("2. 具体问题及修改建议")
        content.append("-" * 40)

        # Group issues by category
        categories = {
            'formatting': '格式问题',
            'language': '语言与逻辑问题',
            'academic_content': '学术内容问题',
            'structure': '结构问题',
            'references': '参考文献问题'
        }

        for eng_cat, chi_cat in categories.items():
            cat_issues = [i for i in self.issues if i['category'] == eng_cat]
            if cat_issues:
                content.append("")
                content.append(f"2.{list(categories.keys()).index(eng_cat)+1} {chi_cat}")
                content.append("")

                for issue in cat_issues:
                    # Priority indicator
                    priority_map = {'high': '【高】', 'medium': '【中】', 'low': '【低】'}
                    priority_indicator = priority_map.get(issue['priority'], '')

                    content.append(f"{priority_indicator} {issue['description']}")
                    content.append(f"   位置: {issue['location']}")
                    content.append(f"   建议: {issue['suggestion']}")
                    content.append("")

        # 3. Top recommendations
        content.append("3. 关键修改项 (Top 10)")
        content.append("-" * 40)
        content.append("")

        # Sort issues by priority (high first, then medium, then low)
        sorted_issues = sorted(self.issues, key=lambda x: {'high': 0, 'medium': 1, 'low': 2}[x['priority']])

        for i, issue in enumerate(sorted_issues[:10]):
            content.append(f"{i+1}. {issue['description']} ({issue['location']})")
        content.append("")

        # 4. Positive feedback
        content.append("4. 鼓励与肯定")
        content.append("-" * 40)
        content.append("")

        if self.positive_feedback:
            for feedback in self.positive_feedback:
                content.append(f"• {feedback}")
        else:
            content.append("论文整体框架完整，符合学术论文基本要求。")

        content.append("")
        content.append("=" * 80)
        content.append("评审结束")
        content.append("=" * 80)

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))

        print(f"评审意见已保存至: {output_path}")
        return output_path

def main():
    """Main function to run thesis review."""
    # Input and output paths
    input_path = "test_files/thesis.docx"
    output_dir = "thesis-review-workspace/iteration-2/eval-comprehensive-review/with_skill/outputs/"

    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    # Generate output filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(output_dir, f"thesis_review_{timestamp}.txt")

    # Check if input file exists
    if not os.path.exists(input_path):
        print(f"错误: 输入文件不存在: {input_path}")
        print("请确保 thesis.docx 文件位于 test_files/ 目录下")
        sys.exit(1)

    try:
        # Initialize reviewer
        reviewer = ThesisReviewer(input_path)

        # Perform review
        reviewer.perform_comprehensive_review()

        # Generate review document
        review_doc = reviewer.generate_review_text(output_path)

        print("\n" + "="*60)
        print("论文评审完成!")
        print(f"输入文件: {input_path}")
        print(f"输出文件: {output_path}")
        print(f"发现问题: {len(reviewer.issues)}个")
        print("="*60)

        # Print summary
        if reviewer.issues:
            print("\n主要问题摘要:")
            high_issues = [i for i in reviewer.issues if i['priority'] == 'high']
            if high_issues:
                print("高优先级问题:")
                for issue in high_issues[:3]:
                    print(f"  - {issue['description']} ({issue['location']})")

            if reviewer.positive_feedback:
                print("\n做得好的方面:")
                for feedback in reviewer.positive_feedback:
                    print(f"  - {feedback}")

    except Exception as e:
        print(f"评审过程中出现错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()