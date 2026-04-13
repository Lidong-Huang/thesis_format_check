#!/usr/bin/env python3
"""
增强版论文评审脚本，包含所有要求的检查项
"""

import re
import os
import sys
import random
from datetime import datetime
from docx import Document

class EnhancedThesisReviewer:
    def __init__(self, input_path):
        """初始化评审器"""
        self.input_path = input_path
        self.document = Document(input_path)
        self.issues = []
        self.positive_feedback = []

        # 提取文档内容
        self.text_content = []
        self.paragraphs = []
        self.headings = []
        self.tables = []
        self.references = []

        self._extract_content()

    def _extract_content(self):
        """从文档中提取所有内容"""
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            if text:
                self.text_content.append(text)
                self.paragraphs.append({
                    'index': i,
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })

                # 检查是否是标题
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    self.headings.append({
                        'index': i,
                        'text': text,
                        'level': level,
                        'para_index': i
                    })

        # 提取表格
        for i, table in enumerate(self.document.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            self.tables.append({
                'index': i,
                'data': table_data,
                'location': f"Table {i + 1}"
            })

    def check_punctuation_consistency(self):
        """检查中英文标点符号混用"""
        chinese_punct = ['。', '，', '；', '：', '？', '！', '「', '」', '《', '》']
        english_punct = ['.', ',', ';', ':', '?', '!', '"', "'", '<', '>']

        for i, para in enumerate(self.paragraphs):
            text = para['text']

            # 检查同一段落中是否混用标点
            has_chinese = any(p in text for p in chinese_punct)
            has_english = any(p in text for p in english_punct[:6])  # 基本标点

            if has_chinese and has_english:
                # 检查是否可能是故意的（英文术语或引用）
                english_terms = ['et al.', 'e.g.', 'i.e.', 'etc.', 'vs.', 'Fig.', 'Table', 'p<', 'p>', 'p=']
                has_english_terms = any(term in text for term in english_terms)

                if not has_english_terms:
                    self.issues.append({
                        'category': 'formatting',
                        'type': 'mixed_punctuation',
                        'location': f"段落 {i+1}",
                        'description': '中英文标点符号混用',
                        'suggestion': '统一使用中文标点符号（。，；：？！）或根据上下文合理选择',
                        'priority': 'medium'
                    })

    def check_chart_caption_clarity(self):
        """检查图表标题是否自明"""
        # 查找图表标题模式
        caption_patterns = [
            r'图\s*\d+(\.\d+)*[\.\s]+(.+)',  # 图 1.1. 标题
            r'Figure\s*\d+(\.\d+)*[\.\s]+(.+)',  # Figure 1.1. 标题
            r'表\s*\d+(\.\d+)*[\.\s]+(.+)',  # 表 1.1. 标题
            r'Table\s*\d+(\.\d+)*[\.\s]+(.+)',  # Table 1.1. 标题
        ]

        for i, para in enumerate(self.paragraphs):
            text = para['text']

            for pattern in caption_patterns:
                match = re.search(pattern, text)
                if match:
                    caption = match.group(2).strip()

                    # 检查标题是否自明
                    # 自明的标题应该包含：什么数据、什么条件、什么结果
                    clarity_indicators = [
                        ('数据', '明确说明数据来源或类型'),
                        ('结果', '说明展示的结果'),
                        ('比较', '说明比较的对象'),
                        ('变化', '说明变化趋势'),
                        ('关系', '说明变量关系')
                    ]

                    missing_indicators = []
                    for indicator, desc in clarity_indicators:
                        if indicator not in caption:
                            missing_indicators.append(desc)

                    if len(missing_indicators) >= 2:  # 缺少多个关键信息
                        self.issues.append({
                            'category': 'academic_content',
                            'type': 'caption_clarity',
                            'location': f"段落 {i+1}: {text[:50]}...",
                            'description': f'图表标题不够自明: "{caption}"',
                            'suggestion': '图表标题应包含数据来源、实验条件、展示结果等关键信息，使读者不看正文也能理解图表内容',
                            'priority': 'medium'
                        })
                    break

    def check_reference_numbering(self):
        """检查文献是否连续编号"""
        ref_section = None
        ref_paras = []

        # 查找参考文献部分
        for i, para in enumerate(self.paragraphs):
            text = para['text']
            if '参考' in text and '文献' in text and len(text) < 20:
                ref_section = i
                break

        if ref_section:
            # 收集参考文献段落
            for para in self.paragraphs[ref_section + 1:]:
                text = para['text']
                if text and not ('Heading' in para['style'] or '致谢' in text or '附录' in text):
                    ref_paras.append(text)
                elif 'Heading' in para['style']:
                    break

            # 检查编号
            expected_num = 1
            for ref_text in ref_paras:
                # 查找 [1], [2] 等编号
                match = re.search(r'\[(\d+)\]', ref_text)
                if match:
                    ref_num = int(match.group(1))
                    if ref_num != expected_num:
                        self.issues.append({
                            'category': 'references',
                            'type': 'reference_numbering',
                            'location': f'参考文献 #{expected_num}',
                            'description': f'参考文献编号不连续: 期望[{expected_num}]，实际[{ref_num}]',
                            'suggestion': '检查并修正参考文献编号顺序',
                            'priority': 'high'
                        })
                        expected_num = ref_num + 1
                    else:
                        expected_num += 1
                else:
                    # 未找到编号，可能使用不同格式
                    break

    def check_reference_authenticity(self):
        """随机抽取5篇文献检查真实性"""
        ref_section = None
        ref_paras = []

        # 查找参考文献部分
        for i, para in enumerate(self.paragraphs):
            text = para['text']
            if '参考' in text and '文献' in text and len(text) < 20:
                ref_section = i
                break

        if ref_section:
            # 收集参考文献段落
            for para in self.paragraphs[ref_section + 1:]:
                text = para['text']
                if text and not ('Heading' in para['style'] or '致谢' in text or '附录' in text):
                    ref_paras.append(text)
                elif 'Heading' in para['style']:
                    break

            if ref_paras:
                # 随机抽取5篇文献（或全部如果少于5篇）
                sample_size = min(5, len(ref_paras))
                if sample_size > 0:
                    sampled_refs = random.sample(ref_paras, sample_size)

                    self.issues.append({
                        'category': 'references',
                        'type': 'reference_verification_needed',
                        'location': '参考文献部分',
                        'description': f'需要验证 {sample_size} 篇参考文献的真实性（随机抽取）',
                        'suggestion': '建议使用学术数据库（如CNKI、Web of Science、Google Scholar）验证以下文献的真实性:\n' +
                                     '\n'.join([f'  {i+1}. {ref[:100]}...' for i, ref in enumerate(sampled_refs)]),
                        'priority': 'medium'
                    })

    def check_abstract_balance(self):
        """检查中文摘要是否结果部分写得太少"""
        abstract_paras = []
        in_abstract = False

        for para in self.paragraphs:
            text = para['text']
            style = para['style']

            if '摘要' in text and ('Heading' in style or len(text) < 20):
                in_abstract = True
                continue
            elif in_abstract:
                if 'Heading' in style:  # 下一个标题结束摘要
                    break
                if text and not text.startswith('关键词'):
                    abstract_paras.append(text)

        if abstract_paras:
            abstract_text = ' '.join(abstract_paras)
            # 简单启发式：统计背景词和结果词
            background_terms = ['随着', '近年来', '目前', '背景', '目的', '旨在', '探讨', '研究', '问题']
            results_terms = ['结果', '发现', '表明', '显示', '达到', '验证', '证明', '结论', '得出']

            bg_count = sum(term in abstract_text for term in background_terms)
            res_count = sum(term in abstract_text for term in results_terms)

            if bg_count > 0 and res_count > 0:
                ratio = bg_count / (bg_count + res_count)
                if ratio > 0.7:  # 超过70%的背景词
                    self.issues.append({
                        'category': 'structure',
                        'type': 'abstract_balance',
                        'location': '摘要部分',
                        'description': f'摘要头重脚轻: 背景介绍占比过高({ratio:.0%})，结果部分不足',
                        'suggestion': '增加结果和结论的篇幅，减少背景介绍',
                        'priority': 'medium'
                    })
                elif res_count == 0:
                    self.issues.append({
                        'category': 'structure',
                        'type': 'abstract_no_results',
                        'location': '摘要部分',
                        'description': '摘要中未提及研究结果',
                        'suggestion': '在摘要中明确说明主要研究结果和结论',
                        'priority': 'high'
                    })

    def check_methods_order(self):
        """检查材料与方法部分是否符合时间顺序"""
        methods_headings = []

        for heading in self.headings:
            text = heading['text']
            if '方法' in text or '实验设计' in text or '材料' in text or '实验' in text:
                methods_headings.append(heading)

        if len(methods_headings) > 1:
            # 检查标题是否按逻辑顺序排列
            # 通常：材料/数据 -> 方法/设计 -> 实施/步骤 -> 分析/评估
            order_keywords = [
                ['材料', '数据', '样本', '来源', '设备'],  # 第一：材料/数据
                ['方法', '设计', '模型', '算法', '方案'],  # 第二：方法/设计
                ['步骤', '流程', '实施', '过程', '操作'],  # 第三：步骤
                ['分析', '评估', '指标', '统计', '测量']   # 第四：分析
            ]

            heading_texts = [h['text'] for h in methods_headings]
            for i, heading in enumerate(heading_texts):
                for stage_idx, keywords in enumerate(order_keywords):
                    if any(keyword in heading for keyword in keywords):
                        if i < stage_idx - 1:  # 太早
                            self.issues.append({
                                'category': 'structure',
                                'type': 'methods_order',
                                'location': f"标题: {heading}",
                                'description': '方法部分顺序可能不符合时间或逻辑顺序',
                                'suggestion': '考虑按"材料/数据 → 方法/设计 → 实施步骤 → 分析评估"的时间或逻辑顺序组织',
                                'priority': 'medium'
                            })
                        break

    def check_title_content_alignment(self):
        """检查大题目和内容是否脱节"""
        if self.paragraphs:
            title = self.paragraphs[0]['text'] if len(self.paragraphs[0]['text']) > 10 else None

            if title and len(title) < 100:  # 可能是主标题
                # 从标题中提取关键词
                title_keywords = re.findall(r'[\u4e00-\u9fff]{2,}', title)

                # 检查这些关键词是否在内容中出现
                content_text = ' '.join([p['text'] for p in self.paragraphs[10:50]])  # 前几段

                missing_keywords = []
                for keyword in title_keywords:
                    if len(keyword) > 1 and keyword not in content_text:
                        missing_keywords.append(keyword)

                if missing_keywords:
                    self.issues.append({
                        'category': 'structure',
                        'type': 'title_content_alignment',
                        'location': '论文标题',
                        'description': f'标题中的关键词"{", ".join(missing_keywords)}"在正文前部出现频率较低',
                        'suggestion': '确保标题准确反映论文核心内容，或在正文中更充分地讨论标题涉及的主题',
                        'priority': 'high'
                    })

    def check_decimal_places(self):
        """检查表格小数位数超过2位的要标出来"""
        for table in self.tables:
            for row_idx, row in enumerate(table['data']):
                for col_idx, cell in enumerate(row):
                    # 查找带小数位的数字
                    numbers = re.findall(r'\d+\.\d+', cell)
                    for num in numbers:
                        decimal_part = num.split('.')[1]
                        if len(decimal_part) > 2:
                            self.issues.append({
                                'category': 'academic_content',
                                'type': 'excessive_decimals',
                                'location': f"{table['location']}, 行{row_idx+1}, 列{col_idx+1}",
                                'description': f'数值"{num}"有{len(decimal_part)}位小数，通常保留2位即可',
                                'suggestion': '除非测量精度要求，否则将小数位数减少到2位',
                                'priority': 'low'
                            })

    def check_figure_table_numbering(self):
        """检查图表编号连续性"""
        fig_patterns = [r'图\s*(\d+(\.\d+)*)', r'Figure\s*(\d+(\.\d+)*)', r'Fig\.\s*(\d+(\.\d+)*)']
        table_patterns = [r'表\s*(\d+(\.\d+)*)', r'Table\s*(\d+(\.\d+)*)', r'Tab\.\s*(\d+(\.\d+)*)']

        fig_numbers = []
        table_numbers = []

        for para in self.paragraphs:
            text = para['text']

            # 检查图引用
            for pattern in fig_patterns:
                matches = re.finditer(pattern, text)
                for match in matches:
                    num = match.group(1)
                    fig_numbers.append(num)

            # 检查表引用
            for pattern in table_patterns:
                matches = re.finditer(pattern, text)
                for match in matches:
                    num = match.group(1)
                    table_numbers.append(num)

        # 检查图编号
        self._check_numbering_sequence(fig_numbers, '图', 'figure_numbering')

        # 检查表编号
        self._check_numbering_sequence(table_numbers, '表', 'table_numbering')

    def _check_numbering_sequence(self, numbers, label, issue_type):
        """检查编号是否连续"""
        if not numbers:
            return

        # 转换为可排序格式
        sorted_nums = sorted(set(numbers), key=lambda x: [int(part) for part in x.split('.')])

        # 检查间隔
        for i in range(1, len(sorted_nums)):
            prev_parts = sorted_nums[i-1].split('.')
            curr_parts = sorted_nums[i].split('.')

            if len(prev_parts) == len(curr_parts):
                try:
                    prev_last = int(prev_parts[-1])
                    curr_last = int(curr_parts[-1])

                    if curr_last != prev_last + 1:
                        self.issues.append({
                            'category': 'academic_content',
                            'type': issue_type,
                            'location': f'{label}编号',
                            'description': f'{label}编号不连续: {sorted_nums[i-1]} 后应为 {prev_last+1}，但找到 {sorted_nums[i]}',
                            'suggestion': f'检查并修正{label}编号顺序',
                            'priority': 'medium'
                        })
                except ValueError:
                    pass

    def collect_positive_feedback(self):
        """收集论文的积极方面"""
        # 检查摘要结构
        abstract_paras = []
        in_abstract = False

        for para in self.paragraphs:
            text = para['text']
            if '摘要' in text and len(text) < 20:
                in_abstract = True
            elif in_abstract and ('Heading' in para['style'] or '关键词' in text):
                break
            elif in_abstract and text:
                abstract_paras.append(text)

        if 2 <= len(abstract_paras) <= 4:
            self.positive_feedback.append('摘要结构清晰，长度适中')

        # 检查标题清晰度
        if len(self.headings) >= 5:
            self.positive_feedback.append('章节结构明确，标题层次清晰')

        # 检查表格/图表
        if self.tables:
            self.positive_feedback.append('包含数据表格，增强论证说服力')

        # 检查参考文献
        ref_count = sum(1 for para in self.paragraphs if '参考' in para['text'] and '文献' in para['text'])
        if ref_count > 0:
            self.positive_feedback.append('包含参考文献部分，体现学术规范')

    def perform_comprehensive_review(self):
        """执行所有检查"""
        print("开始执行增强版论文评审...")

        # 运行所有检查
        checks = [
            ('检查标点符号一致性', self.check_punctuation_consistency),
            ('检查图表标题自明性', self.check_chart_caption_clarity),
            ('检查文献连续编号', self.check_reference_numbering),
            ('检查文献真实性', self.check_reference_authenticity),
            ('检查摘要平衡性', self.check_abstract_balance),
            ('检查方法部分顺序', self.check_methods_order),
            ('检查标题内容一致性', self.check_title_content_alignment),
            ('检查表格小数位数', self.check_decimal_places),
            ('检查图表编号连续性', self.check_figure_table_numbering),
            ('收集积极反馈', self.collect_positive_feedback),
        ]

        for check_name, check_func in checks:
            print(f"  - {check_name}")
            check_func()

        print(f"评审完成，共发现{len(self.issues)}个问题")

    def generate_review_text(self, output_path):
        """生成评审文档"""
        content = []

        # 标题
        content.append("=" * 80)
        content.append("硕士学位论文评审意见（增强版）")
        content.append("=" * 80)
        content.append("")

        # 元数据
        content.append(f"评审文件: {os.path.basename(self.input_path)}")
        content.append(f"生成日期: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        content.append(f"评审人: Claude Code 学术评审系统（增强版）")
        content.append("")

        # 1. 总体评价
        content.append("1. 总体评价")
        content.append("-" * 40)

        # 按优先级统计问题
        high_issues = [i for i in self.issues if i['priority'] == 'high']
        medium_issues = [i for i in self.issues if i['priority'] == 'medium']
        low_issues = [i for i in self.issues if i['priority'] == 'low']

        content.append("本文已按照硕士学位论文要求进行综合评审，涵盖格式规范、学术表达、结构组织和参考文献等方面。")
        content.append(f"共发现{len(self.issues)}个问题，其中：")
        content.append(f"  • 高优先级问题: {len(high_issues)}个")
        content.append(f"  • 中优先级问题: {len(medium_issues)}个")
        content.append(f"  • 低优先级问题: {len(low_issues)}个")
        content.append("")

        # 2. 具体问题及修改建议
        content.append("2. 具体问题及修改建议")
        content.append("-" * 40)

        # 按类别分组问题
        categories = {
            'formatting': '格式问题',
            'academic_content': '学术内容问题',
            'structure': '结构问题',
            'references': '参考文献问题'
        }

        category_order = ['formatting', 'academic_content', 'structure', 'references']

        for eng_cat in category_order:
            chi_cat = categories.get(eng_cat)
            if chi_cat:
                cat_issues = [i for i in self.issues if i['category'] == eng_cat]
                if cat_issues:
                    content.append("")
                    content.append(f"2.{category_order.index(eng_cat)+1} {chi_cat}")
                    content.append("")

                    for issue in cat_issues:
                        # 优先级指示器
                        priority_map = {'high': '【高】', 'medium': '【中】', 'low': '【低】'}
                        priority_indicator = priority_map.get(issue['priority'], '')

                        content.append(f"{priority_indicator} {issue['description']}")
                        content.append(f"   位置: {issue['location']}")
                        content.append(f"   建议: {issue['suggestion']}")
                        content.append("")

        # 3. 关键修改项
        content.append("3. 关键修改项")
        content.append("-" * 40)
        content.append("")

        # 按优先级排序问题
        sorted_issues = sorted(self.issues, key=lambda x: {'high': 0, 'medium': 1, 'low': 2}[x['priority']])

        for i, issue in enumerate(sorted_issues[:10]):
            content.append(f"{i+1}. {issue['description']} ({issue['location']})")
        content.append("")

        # 4. 鼓励与肯定
        content.append("4. 鼓励与肯定")
        content.append("-" * 40)
        content.append("")

        if self.positive_feedback:
            for feedback in self.positive_feedback:
                content.append(f"• {feedback}")
        else:
            content.append("论文整体框架完整，符合学术论文基本要求。")

        content.append("")
        content.append("=" * 80)
        content.append("评审结束")
        content.append("=" * 80)

        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))

        print(f"评审意见已保存至: {output_path}")
        return output_path

def main():
    """主函数"""
    # 输入和输出路径
    input_path = "test_files/论文终稿.docx"
    output_dir = "thesis-review-workspace/iteration-2/eval-focused-check/with_skill/outputs/"

    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)

    # 生成输出文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(output_dir, f"论文终稿_增强评审_{timestamp}.txt")

    # 检查输入文件是否存在
    if not os.path.exists(input_path):
        print(f"错误: 输入文件不存在: {input_path}")
        sys.exit(1)

    try:
        # 初始化评审器
        print(f"开始增强版评审: {input_path}")
        reviewer = EnhancedThesisReviewer(input_path)

        # 执行评审
        reviewer.perform_comprehensive_review()

        # 生成评审文档
        review_doc = reviewer.generate_review_text(output_path)

        print("\n" + "="*60)
        print("增强版论文评审完成!")
        print(f"输入文件: {input_path}")
        print(f"输出文件: {output_path}")
        print(f"发现问题: {len(reviewer.issues)}个")
        print("="*60)

        # 打印摘要
        if reviewer.issues:
            print("\n主要问题摘要:")
            high_issues = [i for i in reviewer.issues if i['priority'] == 'high']
            if high_issues:
                print("高优先级问题:")
                for issue in high_issues[:3]:
                    print(f"  - {issue['description']} ({issue['location']})")

            if reviewer.positive_feedback:
                print("\n做得好的方面:")
                for feedback in reviewer.positive_feedback:
                    print(f"  - {feedback}")

    except Exception as e:
        print(f"评审过程中出现错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()