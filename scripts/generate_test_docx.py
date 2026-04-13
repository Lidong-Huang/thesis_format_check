#!/usr/bin/env python3
"""
Generate a test thesis document with various issues for testing the thesis-review skill.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

def create_test_thesis():
    """Create a test thesis .docx file with intentional issues."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # Title
    title = doc.add_heading('人工智能在医疗诊断中的应用研究', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author info
    doc.add_paragraph('作者：张三')
    doc.add_paragraph('学号：2023123456')
    doc.add_paragraph('指导教师：李四 教授')
    doc.add_paragraph('专业：计算机科学与技术')
    doc.add_paragraph('完成日期：2025年6月')

    doc.add_page_break()

    # Abstract - intentionally head-heavy
    doc.add_heading('摘要', 1)
    abstract = doc.add_paragraph()
    abstract.add_run('随着人工智能技术的飞速发展。')  # Chinese period
    abstract.add_run('AI在医疗领域的应用日益广泛. ')  # Mixed: English period after Chinese text
    abstract.add_run('本文旨在探讨深度学习在医学影像诊断中的有效性。')
    abstract.add_run('通过回顾大量文献，我们发现...')  # Lots of background
    abstract.add_run('本研究采用卷积神经网络对胸部X光片进行分析。')
    abstract.add_run('实验结果表明，模型准确率达到92.3%。')  # Minimal results
    abstract.add_run('该研究为AI辅助诊断提供了参考。')

    doc.add_heading('ABSTRACT', 1)
    doc.add_paragraph('With the rapid development of artificial intelligence technology...')

    doc.add_page_break()

    # Table of contents placeholder
    doc.add_heading('目录', 1)
    doc.add_paragraph('1. 引言 ............... 1')
    doc.add_paragraph('2. 文献综述 ........... 3')
    doc.add_paragraph('3. 研究方法 ........... 5')
    doc.add_paragraph('4. 实验结果与分析 ..... 8')
    doc.add_paragraph('5. 结论与展望 ......... 12')
    doc.add_paragraph('参考文献 ............. 15')
    doc.add_paragraph('致谢 ............... 18')

    doc.add_page_break()

    # Chapter 1
    doc.add_heading('1. 引言', 1)
    intro = doc.add_paragraph()
    intro.add_run('人工智能(Artificial Intelligence, AI)是当前科技发展的热点领域。')
    intro.add_run('在医疗诊断中，AI技术展现出巨大潜力. ')  # Mixed punctuation
    intro.add_run('然而，现有研究仍存在诸多挑战。')

    # Contradiction: paragraph 1 says AI has big potential, paragraph 2 says limited
    doc.add_paragraph('人工智能在医疗诊断中的应用潜力有限，受到数据质量和算法透明度的制约。')  # Contradicts previous

    doc.add_paragraph('本文结构如下：第2章介绍相关文献，第3章阐述研究方法，第4章展示实验结果，第5章总结全文。')

    # Chapter 2 - with reference issues
    doc.add_heading('2. 文献综述', 1)
    doc.add_paragraph('近年来，深度学习在医学影像分析中取得显著进展[1]。')
    doc.add_paragraph('卷积神经网络(CNN)在图像分类任务中表现优异[2]。')
    doc.add_paragraph('然而，模型可解释性仍是挑战[3]。')
    doc.add_paragraph('一些研究尝试结合注意力机制提升性能[4]。')

    # Chapter 3 - Methods with logical order issues
    doc.add_heading('3. 研究方法', 1)
    doc.add_paragraph('本章介绍实验设计、数据收集和模型构建。')  # Logical but not chronological

    doc.add_heading('3.1 数据来源', 2)
    doc.add_paragraph('本研究使用公开胸部X光片数据集。')

    doc.add_heading('3.3 模型评估', 2)  # Skipped 3.2 intentionally
    doc.add_paragraph('采用准确率、召回率和F1分数评估模型性能。')

    doc.add_heading('3.2 模型构建', 2)  # Out of order
    doc.add_paragraph('构建ResNet-50卷积神经网络。')

    # Chapter 4 - Tables and figures
    doc.add_heading('4. 实验结果与分析', 1)

    doc.add_heading('4.1 定量结果', 2)
    doc.add_paragraph('表4.1展示不同模型的性能对比。')

    # Table with >2 decimal places
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Shading'

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '模型'
    hdr_cells[1].text = '准确率'
    hdr_cells[2].text = '召回率'
    hdr_cells[3].text = 'F1分数'

    # Data rows
    data = [
        ('CNN', '0.9234', '0.9123', '0.9178'),
        ('ResNet', '0.9456', '0.9321', '0.9387'),
        ('VGG', '0.9012', '0.8890', '0.8951'),
        ('EfficientNet', '0.9567', '0.9489', '0.9528')
    ]

    for i, (model, acc, rec, f1) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = model
        row_cells[1].text = acc  # Has 4 decimal places
        row_cells[2].text = rec
        row_cells[3].text = f1

    doc.add_paragraph('表4.1 模型性能对比')  # Caption

    doc.add_paragraph('图4.1显示训练损失曲线。')
    doc.add_paragraph('图4.3展示特征可视化结果。')  # Missing 4.2

    doc.add_paragraph('统计分析显示，ResNet性能显著优于CNN(p=0.0432)。')  # p-value with 4 decimal places
    doc.add_paragraph('差异具有统计学意义(p < 0.05)。')  # Correct format

    # Chapter 5
    doc.add_heading('5. 结论与展望', 1)
    doc.add_paragraph('本研究证实了深度学习在医疗诊断中的有效性。')
    doc.add_paragraph('未来工作将探索多模态数据融合。')

    doc.add_page_break()

    # References - with numbering and potential fake references
    doc.add_heading('参考文献', 1)

    refs = [
        'Goodfellow I, Bengio Y, Courville A. Deep Learning. MIT Press, 2016.',
        'Krizhevsky A, Sutskever I, Hinton G E. ImageNet classification with deep convolutional neural networks. NIPS 2012.',
        'Esteva A, Kuprel B, Novoa R A, et al. Dermatologist-level classification of skin cancer with deep neural networks. Nature, 2017, 542(7639): 115-118.',
        'Wang X, Peng Y, Lu L, et al. ChestX-ray8: Hospital-scale chest X-ray database and benchmarks on weakly-supervised classification and localization of common thorax diseases. CVPR 2017.',
        'Smith J, Johnson R. A non-existent paper for testing. Journal of Fake Studies, 2025, 12(3): 45-67.',  # Potentially fake
        'Li H, Wang Z, Liu S. Advanced deep learning for medical imaging. IEEE Transactions on Medical Imaging, 2020, 39(5): 1234-1245.',
        'Zhang Y, Chen Q, Yang J. Explainable AI in healthcare: A survey. Artificial Intelligence in Medicine, 2021, 112: 102038.',
        'Brown T, Mann B, Ryder N, et al. Language models are few-shot learners. NeurIPS 2020.',
        'Fake Author A, Fake Author B. Completely fabricated study. International Journal of Nonexistent Research, 2024, 8(2): 99-120.',  # Potentially fake
        'Zhou B, Khosla A, Lapedriza A, et al. Learning deep features for discriminative localization. CVPR 2016.'
    ]

    for i, ref in enumerate(refs, 1):
        # Intentional skip in numbering
        if i == 6:
            continue  # Skip reference 6 to create numbering gap
        doc.add_paragraph(f'[{i if i < 6 else i-1}] {ref}')

    # Acknowledgments
    doc.add_heading('致谢', 1)
    doc.add_paragraph('感谢导师的指导和同学的支持。')

    # Save
    output_path = 'test_files/thesis_sample.docx'
    doc.save(output_path)
    print(f"Generated test thesis at: {output_path}")

    # Also generate a cleaner version for second test
    doc2 = Document()
    title2 = doc2.add_heading('机器学习在金融风控中的应用研究', 0)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc2.add_paragraph('作者：王五')
    doc2.add_page_break()
    doc2.add_heading('摘要', 1)
    doc2.add_paragraph('本文研究机器学习在金融风险管理中的应用。通过实证分析，我们发现...')
    doc2.add_heading('1. 引言', 1)
    doc2.add_paragraph('金融风险管理是银行业务的核心环节。')
    doc2.save('test_files/thesis.docx')
    print(f"Generated second test thesis at: test_files/thesis.docx")

    doc3 = Document()
    title3 = doc3.add_heading('论文终稿', 0)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc3.add_paragraph('这是论文终稿，包含更多问题供测试。')
    doc3.add_heading('摘要', 1)
    doc3.add_paragraph('背景介绍...背景介绍...背景介绍...背景介绍...背景介绍...结果很少。')
    doc3.add_heading('图表', 1)
    doc3.add_paragraph('图1.1 示意图')
    doc3.add_paragraph('图1.3 另一个图')  # Skip 1.2
    doc3.save('test_files/论文终稿.docx')
    print(f"Generated third test thesis at: test_files/论文终稿.docx")

if __name__ == '__main__':
    try:
        create_test_thesis()
    except ImportError as e:
        print(f"Error: {e}")
        print("Please install python-docx: pip install python-docx")
    except Exception as e:
        print(f"Unexpected error: {e}")